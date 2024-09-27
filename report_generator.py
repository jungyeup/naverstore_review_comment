import os
import logging
from docx import Document
import pandas as pd
from datetime import datetime
from openai import OpenAI

# Ensure logging is configured
logging.basicConfig(level=logging.INFO)

class ReportGenerator:
    def __init__(self, openai_api_key, folder_name='history'):
        self.folder_name = folder_name
        if not os.path.exists(self.folder_name):
            os.makedirs(self.folder_name)

        # Initialize OpenAI client
        self.client = OpenAI(api_key=openai_api_key)
        # Define the system prompt with some base instructions for the GPT model if needed
        self.system_prompt = "한국어로작성해줘"

    def generate_gpt4_summary(self, interaction_data):
        """
        Generate a summary using GPT-4 for the given interaction data.
        """
        try:
            prompt = f"""
            Please provide a detailed description based on the following interaction data:
            {interaction_data}
            """

            chat_completion = self.client.chat.completions.create(
                model="gpt-4o",
                temperature=0.7,
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": prompt}
                ]
            )
            return chat_completion.choices[0].message.content
        except Exception as e:
            logging.error(f"Error generating summary: {e}")
            return f"Error generating summary: {e}"

    def generate_docx_report(self, data, file_path, append=False):
        if append and os.path.exists(file_path):
            doc = Document(file_path)
        else:
            doc = Document()
            doc.add_heading('Daily Report', 0)

        for record in data:
            if not record:
                logging.info("Skipping None record")
                continue  # Skip the record if it is None or empty
            doc.add_heading(record.get('type', 'Unknown'), level=1)
            doc.add_paragraph(f"Timestamp: {record.get('timestamp', 'N/A')}")
            doc.add_paragraph(f"Product Name: {record.get('product_name', 'N/A')}")
            doc.add_paragraph(f"User: {record.get('user', 'N/A')}")
            doc.add_paragraph(f"URL: {record.get('url', 'N/A')}")
            doc.add_paragraph(f"Question: {record.get('question', 'N/A')}")
            doc.add_paragraph(f"Answer: {record.get('answer', 'N/A')}")
            doc.add_paragraph(f"OCR Summaries: {record.get('ocr_summaries', 'N/A')}")
            doc.add_paragraph(f"Status: {record.get('status', 'N/A')}")
            doc.add_paragraph(f"Error Message: {record.get('error_message', 'None')}")

            summary = self.generate_gpt4_summary(record)
            doc.add_paragraph(f"Summary: {summary}")

        doc.save(file_path)
        logging.info(f"Report saved as .docx at {file_path}")

    def generate_xlsx_report(self, data, file_path, append=False):
        if append and os.path.exists(file_path):
            df_existing = pd.read_excel(file_path)
            df_new = pd.DataFrame(data)
            df = pd.concat([df_existing, df_new], ignore_index=True)
        else:
            df = pd.DataFrame(data)
        df.to_excel(file_path, index=False)
        logging.info(f"Report saved as .xlsx at {file_path}")

    def generate_reports(self, data):
        current_date = datetime.now().strftime('%Y-%m-%d')
        docx_file_path = os.path.join(self.folder_name, f'Daily_Report_{current_date}.docx')
        xlsx_file_path = os.path.join(self.folder_name, f'Daily_Report_{current_date}.xlsx')

        logging.info("Generating DOCX report...")
        self.generate_docx_report(data, docx_file_path)

        logging.info("Generating XLSX report...")
        self.generate_xlsx_report(data, xlsx_file_path)